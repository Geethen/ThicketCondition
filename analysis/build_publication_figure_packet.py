"""Build the lead-author figure review packet and the publication asset archive.

The document follows the ``decision_memo`` named override of the
``standard_business_brief`` preset, using the ``memo_masthead`` first-page
header pattern from the document-production skill.
"""

from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FIGURE_DIR = ROOT / "analysis" / "figures"
DOCX_PATH = FIGURE_DIR / "publication_figure_review_packet.docx"
ZIP_PATH = FIGURE_DIR / "publication_figure_assets.zip"

BLUE = RGBColor(31, 78, 121)
INK = RGBColor(32, 36, 38)
MUTED = RGBColor(90, 98, 104)
RULE = "C6CFD4"
PALE_GREEN = "EDF3F0"
GREEN = "5C7F71"


def set_cell_border_unused() -> None:
    """Layout tables are intentionally not used in this document."""


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.extend([begin, instruction, separate, value, end])


def add_bottom_border(paragraph, color: str = RULE, size: str = "10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def shade_callout(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), PALE_GREEN)
    p_pr.append(shading)

    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), GREEN)
    borders.append(left)
    p_pr.append(borders)

    spacing = p_pr.get_or_add_spacing()
    spacing.set(qn("w:before"), "100")
    spacing.set(qn("w:after"), "140")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.widow_control = True

    heading1 = styles["Heading 1"]
    heading1.font.name = "Arial"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = BLUE
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.keep_with_next = True
    heading1.paragraph_format.keep_together = True

    caption = styles.add_style("Publication Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Arial"
    caption.font.size = Pt(9)
    caption.font.color.rgb = INK
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.space_before = Pt(5)
    caption.paragraph_format.space_after = Pt(0)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_together = True
    caption.paragraph_format.widow_control = True

    document.core_properties.title = "Publication figure review packet"
    document.core_properties.subject = "Figures and captions for lead-author review"
    document.core_properties.author = "Thicket condition mapping team"
    document.core_properties.keywords = "publication figures; area estimates; accuracy; symbolic regression"

    # Populate every Word header/footer variant. This keeps page geometry stable
    # even when a local Word template has odd/even or first-page options enabled.
    section.different_first_page_header_footer = False
    document.settings.odd_and_even_pages_header_footer = True
    for header in (section.header, section.even_page_header, section.first_page_header):
        header_p = header.paragraphs[0]
        header_p.text = "THICKET CONDITION MAPPING  |  LEAD-AUTHOR REVIEW"
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_p.paragraph_format.space_after = Pt(0)
        for run in header_p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = MUTED

    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        footer_p = footer.paragraphs[0]
        footer_p.clear()
        add_page_number(footer_p)


def add_masthead(document: Document) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("PUBLICATION FIGURE PACKAGE")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = BLUE

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Thicket condition mapping paper")
    run.font.name = "Arial"
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = INK

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(9)
    subtitle.paragraph_format.keep_with_next = True
    run = subtitle.add_run("Three candidate figures and publication-ready captions")
    run.font.name = "Arial"
    run.font.size = Pt(13.5)
    run.font.color.rgb = MUTED

    metadata = [
        ("Prepared for", "Lead author"),
        ("Purpose", "Content, framing and caption review"),
        ("Prepared", "6 August 2026"),
    ]
    last_p = None
    for label, value in metadata:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        label_run = p.add_run(f"{label}: ")
        label_run.font.name = "Arial"
        label_run.font.size = Pt(10)
        label_run.font.bold = True
        label_run.font.color.rgb = MUTED
        value_run = p.add_run(value)
        value_run.font.name = "Arial"
        value_run.font.size = Pt(10)
        value_run.font.color.rgb = INK
        last_p = p
    if last_p is not None:
        last_p.paragraph_format.space_after = Pt(5)
        add_bottom_border(last_p)

    callout = document.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.keep_together = True
    shade_callout(callout)
    lead = callout.add_run("Review focus. ")
    lead.bold = True
    lead.font.name = "Arial"
    lead.font.size = Pt(9.5)
    body = callout.add_run(
        "Please check the interpretation of the design-based area estimates, the "
        "reference-label selection rule and the validation caveat in Figure 3. The "
        "archive contains high-resolution PNG and editable SVG versions of every figure."
    )
    body.font.name = "Arial"
    body.font.size = Pt(9.5)


def add_math_symbol(paragraph, base: str, subscript: str | None = None) -> None:
    run = paragraph.add_run(base)
    run.italic = True
    if subscript:
        sub = paragraph.add_run(subscript)
        sub.font.subscript = True


def add_figure(document: Document, number: int, short_title: str, image_name: str, alt_text: str, width: float):
    if number > 1:
        document.add_page_break()
    heading = document.add_paragraph(f"Figure {number}  |  {short_title}", style="Heading 1")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(FIGURE_DIR / image_name), width=Inches(width))._inline
    inline.docPr.set("descr", alt_text)
    inline.docPr.set("title", f"Figure {number}")
    return document.add_paragraph(style="Publication Caption")


def caption_figure_1(paragraph) -> None:
    title = paragraph.add_run("Figure 1 | Design-based area estimates for thicket condition and ecosystem type. ")
    title.bold = True
    panel = paragraph.add_run("a, ")
    panel.bold = True
    paragraph.add_run(
        "Total mapped area in each reference-condition class; parenthetical values give the "
        "corresponding 95% confidence interval as a percentage of the mapped domain. "
    )
    panel = paragraph.add_run("b, ")
    panel.bold = True
    paragraph.add_run(
        "Reference-condition area within Arid, Valley and Mesic thicket. Points show stratified "
        "design-based estimates and horizontal bars show 95% confidence intervals calculated "
        "following the Olofsson estimator. At all locations with two determinate labels (n = 92), one "
        "submitted source label was selected independently with equal probability using NumPy PCG64 "
        "(seed 0); unsure labels were ineligible, and locations with no determinate label (n = 25) were "
        "excluded. Reference observations labelled as no thicket were grouped with the severe class. The "
        "reference sample comprised 821 independently labelled points (Arid, "
    )
    n_run = paragraph.add_run("n")
    n_run.italic = True
    paragraph.add_run(" = 315; Valley, ")
    n_run = paragraph.add_run("n")
    n_run.italic = True
    paragraph.add_run(" = 385; Mesic, ")
    n_run = paragraph.add_run("n")
    n_run.italic = True
    paragraph.add_run(" = 121) across a mapped domain of 1,898,600 ha.")


def caption_figure_2(paragraph) -> None:
    title = paragraph.add_run("Figure 2 | Design-based accuracy of the three-class thicket-condition map. ")
    title.bold = True
    paragraph.add_run(
        "Rows represent mapped classes and columns represent reference classes. Confusion-matrix "
        "cells and "
    )
    n_run = paragraph.add_run("n")
    n_run.italic = True
    paragraph.add_run(
        " values are unweighted reference-point counts; user's, producer's and overall accuracies "
        "are area-adjusted design-based estimates. Accuracy values are percentages, with 95% "
        "confidence intervals in parentheses. At all locations with two determinate labels (n = 92), one "
        "submitted source label was selected independently with equal probability using NumPy PCG64 "
        "(seed 0); unsure labels were ineligible, and locations with no determinate label (n = 25) were "
        "excluded. Reference observations labelled as no thicket were grouped with the severe class. "
        "Estimates use 821 independently labelled reference points."
    )


def caption_figure_3(paragraph) -> None:
    title = paragraph.add_run("Figure 3 | Held-out evaluation of single- and multi-layer intact-class scores. ")
    title.bold = True
    panel = paragraph.add_run("a, ")
    panel.bold = True
    paragraph.add_run("Difference in held-out intact-class ")
    add_math_symbol(paragraph, "F", "1")
    paragraph.add_run(" between each fixed multi-layer score and the ")
    add_math_symbol(paragraph, "p", "intact")
    paragraph.add_run(
        " baseline. Points show observed differences and horizontal bars show percentile 95% "
        "confidence intervals from 10,000 paired bootstrap resamples of the 27 held-out 0.2° "
        "spatial blocks ("
    )
    n_run = paragraph.add_run("n")
    n_run.italic = True
    paragraph.add_run(
        " = 704 observations). The vertical line denotes no difference; all displayed intervals "
        "include zero. Rule thresholds were selected using the training folds only ("
    )
    n_run = paragraph.add_run("n")
    n_run.italic = True
    paragraph.add_run(" = 1,379). ")
    panel = paragraph.add_run("b, ")
    panel.bold = True
    paragraph.add_run("Held-out ")
    add_math_symbol(paragraph, "F", "1")
    paragraph.add_run(
        " for every expression on the PySR training-loss/complexity front. Circles denote "
        "expressions containing "
    )
    add_math_symbol(paragraph, "p", "i")
    paragraph.add_run(
        " only and triangles denote expressions that additionally contain "
    )
    add_math_symbol(paragraph, "p", "m")
    paragraph.add_run(" or ")
    add_math_symbol(paragraph, "p", "s")
    paragraph.add_run("; the horizontal line is the ")
    add_math_symbol(paragraph, "p", "i")
    paragraph.add_run(" baseline. Here ")
    add_math_symbol(paragraph, "p", "i")
    paragraph.add_run(", ")
    add_math_symbol(paragraph, "p", "m")
    paragraph.add_run(" and ")
    add_math_symbol(paragraph, "p", "s")
    paragraph.add_run(
        " are the intact, moderate and severe Random Forest probability layers. The bootstrap "
        "comparison is conditional on the cached five-fold out-of-fold probability predictions "
        "and is not a fully nested end-to-end validation of the complete model-selection pipeline."
    )


def build_document() -> None:
    document = Document()
    configure_document(document)
    add_masthead(document)

    p = add_figure(
        document,
        1,
        "Design-based area estimates",
        "figure1_area_estimates.png",
        "Two-panel figure showing design-based thicket-condition area estimates overall and by ecosystem type, with 95 percent confidence intervals.",
        6.45,
    )
    caption_figure_1(p)

    p = add_figure(
        document,
        2,
        "Map accuracy",
        "figure2_accuracy_table.png",
        "Accuracy table showing confusion-matrix counts and area-adjusted user's, producer's and overall accuracy with 95 percent confidence intervals.",
        5.75,
    )
    caption_figure_2(p)

    p = add_figure(
        document,
        3,
        "Single- versus multi-layer intact-class scores",
        "figure3_symbolic_regression.png",
        "Two-panel figure comparing held-out intact-class performance for fixed multi-layer scores and symbolic-regression expressions against the intact-probability baseline.",
        6.45,
    )
    caption_figure_3(p)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)


def build_archive() -> None:
    names = [
        "figure1_area_estimates.png",
        "figure1_area_estimates.svg",
        "figure2_accuracy_table.png",
        "figure2_accuracy_table.svg",
        "figure3_symbolic_regression.png",
        "figure3_symbolic_regression.svg",
    ]
    with ZipFile(ZIP_PATH, "w", compression=ZIP_DEFLATED, compresslevel=9) as archive:
        for name in names:
            archive.write(FIGURE_DIR / name, arcname=name)


if __name__ == "__main__":
    build_document()
    build_archive()
    print(DOCX_PATH)
    print(ZIP_PATH)
